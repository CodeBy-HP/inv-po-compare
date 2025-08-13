import pandas as pd
import streamlit as st
from docx import Document
import json
from typing import Dict, Any, List

def extract_word_info(uploaded_file) -> Dict[str, Any]:
    """
    Extract structured information from Word document for LLM processing
    """
    try:
        # Read the Word document
        doc = Document(uploaded_file)
        
        structured_info = {
            "file_name": uploaded_file.name,
            "document_type": "Word",
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "content": {
                "paragraphs": [],
                "tables": [],
                "headings": [],
                "structured_data": []
            },
            "metadata": {
                "has_tables": len(doc.tables) > 0,
                "has_images": False,  # Will be enhanced later
                "estimated_pages": max(1, len(doc.paragraphs) // 20)  # Rough estimate
            }
        }
        
        # Extract paragraphs and identify headings
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:  # Skip empty paragraphs
                para_info = {
                    "paragraph_number": i + 1,
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "Normal",
                    "is_heading": paragraph.style.name.startswith('Heading') if paragraph.style else False
                }
                
                structured_info["content"]["paragraphs"].append(para_info)
                
                # Collect headings separately
                if para_info["is_heading"]:
                    structured_info["content"]["headings"].append({
                        "level": paragraph.style.name,
                        "text": text,
                        "paragraph_number": i + 1
                    })
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            headers = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                
                if row_idx == 0:  # Assume first row is headers
                    headers = row_data
                else:
                    table_data.append(row_data)
            
            table_info = {
                "table_number": table_idx + 1,
                "headers": headers,
                "rows": len(table_data),
                "columns": len(headers),
                "data": table_data[:10],  # Show only first 10 rows to avoid huge output
                "total_rows": len(table_data)
            }
            
            structured_info["content"]["tables"].append(table_info)
        
        # Try to identify structured data patterns
        structured_data = _identify_business_data_patterns(structured_info["content"])
        structured_info["content"]["structured_data"] = structured_data
        
        # Print structured information in a clean JSON format
        print("\n" + "="*80)
        print("WORD DOCUMENT ANALYSIS - STRUCTURED FORMAT FOR LLM")
        print("="*80)
        print(json.dumps(structured_info, indent=2, ensure_ascii=False))
        print("="*80 + "\n")
        
        return structured_info
        
    except Exception as e:
        error_info = {
            "error": True,
            "message": str(e),
            "file_name": uploaded_file.name if uploaded_file else "Unknown",
            "document_type": "Word"
        }
        print(f"\n[ERROR] Word extraction failed:")
        print(json.dumps(error_info, indent=2))
        return error_info

def _identify_business_data_patterns(content: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Identify potential business data patterns in the Word document
    """
    patterns = []
    
    # Look for invoice/order patterns in text
    invoice_keywords = ["invoice", "bill", "payment", "amount", "total", "due", "order"]
    customer_keywords = ["customer", "client", "vendor", "supplier", "company"]
    
    for paragraph in content["paragraphs"]:
        text_lower = paragraph["text"].lower()
        
        # Check for invoice-like content
        if any(keyword in text_lower for keyword in invoice_keywords):
            patterns.append({
                "pattern_type": "invoice_related",
                "confidence": "medium",
                "location": f"paragraph_{paragraph['paragraph_number']}",
                "text_snippet": paragraph["text"][:100] + "..." if len(paragraph["text"]) > 100 else paragraph["text"]
            })
        
        # Check for customer/vendor information
        if any(keyword in text_lower for keyword in customer_keywords):
            patterns.append({
                "pattern_type": "entity_information",
                "confidence": "medium", 
                "location": f"paragraph_{paragraph['paragraph_number']}",
                "text_snippet": paragraph["text"][:100] + "..." if len(paragraph["text"]) > 100 else paragraph["text"]
            })
    
    # Analyze tables for business data
    for table in content["tables"]:
        if table["headers"]:
            header_text = " ".join(table["headers"]).lower()
            
            if any(keyword in header_text for keyword in invoice_keywords):
                patterns.append({
                    "pattern_type": "invoice_table",
                    "confidence": "high",
                    "location": f"table_{table['table_number']}",
                    "headers": table["headers"]
                })
            elif any(keyword in header_text for keyword in ["product", "item", "description", "quantity", "price"]):
                patterns.append({
                    "pattern_type": "product_catalog",
                    "confidence": "high",
                    "location": f"table_{table['table_number']}",
                    "headers": table["headers"]
                })
    
    return patterns
